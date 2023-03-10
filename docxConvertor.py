from docx.document import Document
from docx.table import _Cell as DocumentCell
from docx.text.paragraph import Paragraph as DocumentParagraph
from docx.styles.styles import Styles
from docx.styles.style import _ParagraphStyle, _CharacterStyle

import markdown
from markdown import MarkdownConfig
from utils import gen_placeholder


class DOCXConvertor:
    document: Document
    styles: Styles
    md: markdown.Markdown

    def __init__(self, base_document="default.docx"):
        import docx
        self.document = docx.Document(base_document)
        self.styles = self.document.styles
        self.md = markdown.Markdown()

    def read(self, doc: str):
        self.md.parse(doc)
        self.set_style(self.md.markdownConfig)
        self.gen(self.md.document, self.document)

    def save(self, path):
        self.document.save(path)

    def set_style(self, o: MarkdownConfig):
        from docx.enum.style import WD_STYLE_TYPE
        for _, v in o.items():
            if v.docx_style_name != "":
                if isinstance(v, MarkdownConfig.Base):
                    if v.docx_style_name in self.styles:
                        cur_style = self.styles[v.docx_style_name]
                    else:
                        cur_style = self.styles.add_style(v.docx_style_name, WD_STYLE_TYPE.PARAGRAPH)
                        cur_style.name = v.docx_style_name
                    if isinstance(v, MarkdownConfig.Heading):
                        cur_style.base_style = self.styles["Heading"]
                    else:
                        cur_style.base_style = self.styles["Normal"]
                    self.set_paragraph_style(v, cur_style)
                    self.set_character_style(v, cur_style)
                elif isinstance(v, MarkdownConfig.FontBase):
                    if v.docx_style_name in self.styles:
                        cur_style = self.styles[v.docx_style_name]
                    else:
                        cur_style = self.styles.add_style(v.docx_style_name, WD_STYLE_TYPE.CHARACTER)
                        cur_style.name = v.docx_style_name
                    self.set_character_style(v, cur_style)

    def gen(self, c: markdown.Component, p: any):
        getattr(self, c.__class__.__name__)(c, p)

    def set_character_style(self, rule: MarkdownConfig.FontBase, style: _CharacterStyle):
        from docx.shared import Pt
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from utils import convert_color
        from docx.enum.text import WD_COLOR_INDEX
        for item in rule.dict(exclude_none=True).items():
            match item:
                case ["en_font", font]:
                    style.font.name = font
                case ["cn_font", font]:
                    style.element.rPr.rFonts.set(qn('w:eastAsia'), font)
                case ["color", color]:
                    style.font.color.rgb = RGBColor.from_string(convert_color(color).lstrip("#"))
                case ["font_size", size]:
                    style.font.size = Pt(size)
                case ["bold", bold]:
                    style.font.bold = bold
                case ["italic", italic]:
                    style.font.italic = italic
                case ["background", background]:
                    background = background.upper()
                    if hasattr(WD_COLOR_INDEX, background):
                        style.font.highlight_color = getattr(WD_COLOR_INDEX, background)
                    else:
                        style.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def set_paragraph_style(self, rule: MarkdownConfig.Base, style: _ParagraphStyle | _CharacterStyle):
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for item in rule.dict(exclude_none=True).items():
            match item:
                case ["first_line_indent", indent]:
                    if rule.font_size is None:
                        style.paragraph_format.first_line_indent = style.base_style.font.size * indent
                    else:
                        style.paragraph_format.first_line_indent = rule.font_size * indent
                case ["line_spacing_type", "1"]:
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                case ["line_spacing_type", "1.5"]:
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                case ["line_spacing_type", "2"]:
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                case ["line_spacing", spacing]:
                    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    style.paragraph_format.line_spacing = Pt(spacing)
                case ["alignment", alignment] | ["block_alignment", alignment]:
                    style.paragraph_format.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, alignment.upper())
                case ["word_wrap", True]:
                    style.paragraph_format.first_line_indent = 0
                    wordWrap = OxmlElement("w:wordWrap")
                    wordWrap.set(qn('w:val'), "0")
                    style.paragraph_format.element.get_or_add_pPr().append(wordWrap)

    def Section(self, o: markdown.Section, p: Document):
        for c in o.components:
            self.gen(c, p)

    def Title(self, o: markdown.Title, p: Document):
        placeholders = gen_placeholder(o)
        p.add_heading(o.get_config().format.format(**placeholders), o.level)

    def Paragraph(self, o: markdown.Paragraph, p: Document):
        paragraph = p.add_paragraph(None, o.get_config().docx_style_name)
        for c in o.components:
            self.gen(c, paragraph)

    def Image(self, o: markdown.Image, p: Document):
        from docx.shared import Pt
        h = Pt(o.get_config().height) if o.get_config().height is not None else None
        w = Pt(o.get_config().width) if o.get_config().width is not None else None
        if h is None and w is None:
            w = p.sections[0].page_width - p.sections[0].right_margin - p.sections[0].left_margin
        p.add_picture(o.link, w, h)
        placeholders = gen_placeholder(o)
        p.add_paragraph(o.get_config().format.format(**placeholders), "Image Label")

    def Code(self, o: markdown.Code, p: Document):
        table = p.add_table(1, 1, "Table Grid")
        table.cell(0, 0).text = o.code
        placeholders = gen_placeholder(o)
        p.add_paragraph(o.get_config().format.format(**placeholders), "Code Label")

    def Table(self, o: markdown.Table, p: Document):
        placeholders = gen_placeholder(o)
        p.add_paragraph(o.get_config().format.format(**placeholders), "Table Label")
        table = p.add_table(o.row_num, o.col_num, "Table Grid")
        for row in range(len(o)):
            for col in range(len(o.components[row])):
                cell = table.cell(row, col)
                self.gen(o[row][col], cell)

    def TableHeadCell(self, o: markdown.TableHeadCell, p: DocumentCell):
        p = p.paragraphs[0]
        p.style = o.get_config().docx_style_name
        for c in o.components:
            self.gen(c, p)

    def TableCell(self, o: markdown.TableHeadCell, p: DocumentCell):
        p = p.paragraphs[0]
        p.style = o.get_config().docx_style_name
        for c in o.components:
            self.gen(c, p)

    def Text(self, o: markdown.Text, p: DocumentParagraph):
        p.add_run(o.text)

    def ImageRef(self, o: markdown.ImageRef, p: DocumentParagraph):
        image = o.markdown.image_dict[o.name]
        placeholders = gen_placeholder(image)
        p.add_run(o.get_config().format.format(**placeholders), o.get_config().docx_style_name)

    def TableRef(self, o: markdown.ImageRef, p: DocumentParagraph):
        table = o.markdown.table_dict[o.name]
        placeholders = gen_placeholder(table)
        p.add_run(o.get_config().format.format(**placeholders), o.get_config().docx_style_name)

    def CodeRef(self, o: markdown.CodeRef, p: DocumentParagraph):
        code = o.markdown.code_dict[o.name]
        placeholders = gen_placeholder(code)
        p.add_run(o.get_config().format.format(**placeholders), o.get_config().docx_style_name)

    def Ref(self, o: markdown.Ref, p: DocumentParagraph):
        p.add_run(str(o.number)).font.superscript = True

    def ItalicSpan(self, o: markdown.ItalicSpan, p: DocumentParagraph):
        p.add_run(o.text).font.italic = True

    def BoldSpan(self, o: markdown.ItalicSpan, p: DocumentParagraph):
        p.add_run(o.text).font.bold = True

    def EscapeSpan(self, o: markdown.EscapeSpan, p: DocumentParagraph):
        p.add_run(o.text)

    def CodeSpan(self, o: markdown.CodeSpan, p: DocumentParagraph):
        p.add_run(o.text, o.get_config().docx_style_name)
